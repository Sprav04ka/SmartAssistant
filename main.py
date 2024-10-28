import openai
import requests
from docx import Document


# Получение списка новостей по запросу
def get_news(query, api_key):
    url = "https://newsapi.org/v2/everything"
    
    # Параметры для API запроса
    params = {
        "q": query,
        "language": "ru",
        "sortBy": "publishedAt",
        "pageSize": 10,
        "apiKey": api_key,
    }

    response = requests.get(url, params=params)
    print(f"Статус ответа: {response.status_code}")

    if response.status_code == 200:
        articles = response.json().get("articles", [])
        print(f"Найдено {len(articles)} статей")
        return articles
    else:
        print(f"Ошибка запроса: {response.status_code}")
        return []

# Фильтрация новостей по ключевым словам, связанным с образованием детей
def filter_news_for_moms(articles):
    keywords = [
        "мама", "дети", "образование", "обучение", "дошкольники",
        "школьники", "навыки", "математика", "чтение", "развитие",
        "воспитание", "школа", "детский сад", "раннее развитие", "учеба",
        "детский", "школьный"
    ]

    filtered_articles = []
    for article in articles:
        for keyword in keywords:
            if keyword.lower() in article["title"].lower() or keyword.lower() in article["description"].lower():
                print(f"Статья '{article['title']}' содержит ключевое слово: {keyword}")
                filtered_articles.append(article)
                break
    print(f"Найдено {len(filtered_articles)} отфильтрованных статей")
    return filtered_articles

# Генерация текста для блога на основе описания статьи
def generate_blog_text(article_title, article_description, api_key_openai):
    openai.api_key = api_key_openai
    prompt = (
        f"Напиши короткий, емкий текст для блога на тему '{article_title}'. "
        f"Целевая аудитория — родители и педагоги, интересующиеся детским образованием. "
        f"Используй простой, информативный стиль, чтобы рассказать о важности темы: {article_description}"
    )

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "system",
                "content": "Ты — эксперт по детскому обучению, и пишешь статьи для блога для мам.",
            },
            {"role": "user", "content": prompt},
        ],
        max_tokens=500,
        temperature=0.7,
    )

    return response["choices"][0]["message"]["content"].strip()

# Генерация текста для баннера по теме статьи
def generate_banner_text(article_title, api_key_openai):
    prompt = (
        f"Создай короткий и запоминающийся текст для инфографики, баннера или визуального контента, "
        f"предназначенного для родителей и педагогов. Основная идея: '{article_title}'. "
        f"Текст должен быть ёмким и подходить для иллюстрации."
    )

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Ты — эксперт по детскому образованию, пишущий для родителей и педагогов."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=100,
        temperature=0.7,
    )

    return response["choices"][0]["message"]["content"].strip()

# Создание изображения для баннера с помощью DALL-E
def generate_image_from_text(banner_text, i, api_key_openai):
    openai.api_key = api_key_openai
    prompt = (
        f"Create a colorful picture for a blog in realism style on the topic: '{banner_text}'. "
        f"The image should include elements related to children's learning and be in a bright and fun style that parents and educators can understand. Use visual images such as books, school supplies, children studying."
    )

    response = openai.Image.create(
        prompt=prompt,
        n=1,
        size="1024x1024",
    )

    # Сохраняем изображение
    image_url = response['data'][0]['url']
    image_data = requests.get(image_url).content
    filename = f"generated_image_{i}.png"  # Динамическое имя файла с индексом i
    with open(filename, 'wb') as file:
        file.write(image_data)
    
    print(f"Изображение сохранено как {filename}")

# Сохранение текста в формате DOCX и TXT
def save_texts(title, blog_content, banner_content, blog_filename, banner_filename):
    # Сохранение текста блога в DOCX
    document = Document()
    document.add_heading(title, 0)
    document.add_paragraph(blog_content)
    document.save(blog_filename)
    print(f"Текст блога сохранён в файл {blog_filename}")

    # Сохранение текста баннера в TXT
    with open(banner_filename, "w", encoding="utf-8") as txt_file:
        txt_file.write(banner_content)
    print(f"Текст баннера сохранён в файл {banner_filename}")

# Основной код выполнения
if __name__ == "__main__":
    api_key_news = "APIKEY"  # API-ключ для NewsAPI
    api_key_openai = "APIKEY"  # API-ключ для OpenAI
    query = (
    "образование детей OR обучение детей OR развитие детей OR "
    "детские навыки OR школьное обучение OR дошкольники OR "
    "дошкольное развитие OR математика для детей OR чтение для детей"
)

    # Получаем новости по запросу
    articles = get_news(query, api_key_news)

    # Проверяем, есть ли отфильтрованные статьи
    if articles:
        filtered_articles = filter_news_for_moms(articles)

        if filtered_articles:
            for i, article in enumerate(filtered_articles, start=1):
                title = article["title"]
                description = article["description"]
                url = article["url"]

                print(f"Генерация текста для статьи: {title}")

                # Генерация текста для блога и баннера
                blog_text = generate_blog_text(title, description, api_key_openai)
                banner_text = generate_banner_text(blog_text, api_key_openai)

                # Сохраняем тексты и изображение
                blog_filename = f"blog_post_{i}.docx"
                banner_filename = f"banner_text_{i}.txt"
                save_texts(title, blog_text, banner_text, blog_filename, banner_filename)

                # Создание изображения для баннера
                img_filename = f"banner_image_{i}.png"
                generate_image_from_text(banner_text, i, api_key_openai)

        else:
            print("Нет подходящих статей для генерации текста.")
    else:
        print("API не вернуло статьи.")
